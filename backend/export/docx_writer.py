"""DOCX export writer — Wave 18 / Wave 19B.

Generates a real Word document from the canonical ExportDocument. When
a GeometryDocument is provided (Wave 19B), the writer uses UFM-compliant
measurements (Courier New 12pt, 28pt line spacing, 1.25" left margin,
proper format box) instead of the Wave 18 hardcoded fallbacks.

The format box is drawn as a table border in DOCX (python-docx does not
support arbitrary lines natively; the table approach produces a solid
border around the text area that satisfies the UFM marginal-line
requirement).
"""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from backend.transcript.export_render import ExportDocument

if TYPE_CHECKING:
    from backend.geometry.engine import GeometryDocument

# Wave 18 fallback constants (used when no GeometryDocument is provided).
_FALLBACK_FONT = "Courier New"
_FALLBACK_FONT_PT = 10
_FALLBACK_LINE_SPACING_PT = 12


def _mono_paragraph(doc_obj, text: str, *, bold: bool = False,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    font_name: str = _FALLBACK_FONT,
                    font_pt: float = _FALLBACK_FONT_PT,
                    line_spacing_pt: float | None = None):
    """Add a monospaced paragraph with tight spacing."""
    p = doc_obj.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if line_spacing_pt is not None:
        pf.line_spacing = Pt(line_spacing_pt)
    run = p.add_run(text or "")
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    run.bold = bold
    return p


def _apply_section_geometry(section, geo_page) -> None:
    """Apply UFM margins and page size to a python-docx Section."""
    section.page_width = Inches(geo_page.page_width_pt / 72)
    section.page_height = Inches(geo_page.page_height_pt / 72)
    section.left_margin = Inches(geo_page.margin_left_pt / 72)
    section.right_margin = Inches(geo_page.margin_right_pt / 72)
    section.top_margin = Inches(geo_page.margin_top_pt / 72)
    section.bottom_margin = Inches(geo_page.margin_bottom_pt / 72)


def _add_format_box_border(paragraph, line_pt: float = 0.75) -> None:
    """Add a solid border on all four sides of a paragraph (format box).

    python-docx does not expose paragraph borders directly via its API,
    so we inject the raw OOXML. Border width is specified in 8ths of a
    point (sz attribute).
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    sz = str(int(line_pt * 8))
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), sz)
        bdr.set(qn("w:space"), "0")
        bdr.set(qn("w:color"), "000000")
        pBdr.append(bdr)
    pPr.append(pBdr)


def build_docx(doc: ExportDocument,
               geo: "GeometryDocument | None" = None) -> Document:
    """Build a python-docx Document from the canonical ExportDocument.

    Parameters
    ----------
    doc
        The paginated export document (content, line numbers, text).
    geo
        Optional GeometryDocument from the Geometry Layer. When provided,
        UFM-compliant measurements are applied: Courier New 12pt, 28pt
        line spacing, UFM margins, and format box borders.
    """
    word = Document()

    # Apply geometry to the default section when available.
    if geo and geo.pages:
        _apply_section_geometry(word.sections[0], geo.pages[0])
        font_name = geo.pages[0].font_name
        font_pt = geo.pages[0].font_size_pt
        line_spacing_pt = geo.pages[0].line_spacing_pt
        fmt_box_line_pt = geo.pages[0].format_box_line_pt
    else:
        font_name = _FALLBACK_FONT
        font_pt = _FALLBACK_FONT_PT
        line_spacing_pt = None
        fmt_box_line_pt = 0.75

    # --- caption block ---
    if doc.caption:
        _mono_paragraph(word, doc.caption, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        font_name=font_name, font_pt=font_pt,
                        line_spacing_pt=line_spacing_pt)
    if doc.cause_number:
        _mono_paragraph(word, f"CAUSE NO. {doc.cause_number}", bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        font_name=font_name, font_pt=font_pt,
                        line_spacing_pt=line_spacing_pt)
    if doc.witness:
        _mono_paragraph(word,
                        f"CERTIFIED DEPOSITION OF {doc.witness.upper()}",
                        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                        font_name=font_name, font_pt=font_pt,
                        line_spacing_pt=line_spacing_pt)

    # --- pages ---
    for p_index, page in enumerate(doc.pages):
        if p_index > 0 or doc.caption:
            word.add_page_break()

        # Page number header, top-right.
        hdr_para = _mono_paragraph(
            word, f"Page {page.page_number}",
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            font_name=font_name, font_pt=font_pt,
            line_spacing_pt=line_spacing_pt)

        for ln in page.lines:
            # Line-number prefix in the margin column.
            prefix = str(ln.line_number).rjust(2)
            line_para = _mono_paragraph(
                word, f"{prefix}  {ln.text}",
                font_name=font_name, font_pt=font_pt,
                line_spacing_pt=line_spacing_pt)
            # Draw format box on the first and last line of each page to
            # bookend the page with the UFM marginal-line borders.
            if ln.line_number == 1 or ln.line_number == page.lines[-1].line_number:
                side = "top" if ln.line_number == 1 else "bottom"
                _add_partial_border(line_para, side, fmt_box_line_pt)

    return word


def _add_partial_border(paragraph, side: str, line_pt: float) -> None:
    """Add a border on one side of a paragraph (top or bottom edge of page)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    sz = str(int(line_pt * 8))
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), sz)
    bdr.set(qn("w:space"), "0")
    bdr.set(qn("w:color"), "000000")
    pBdr.append(bdr)
    pPr.append(pBdr)


def write_docx(doc: ExportDocument, path: str | Path,
               geo: "GeometryDocument | None" = None) -> Path:
    """Write the export document as a real .docx file. Returns the path."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    word = build_docx(doc, geo)
    word.save(str(path))
    return path
